"""
Shared export service — renders a simple intermediate document (ExportDoc) to
PDF / DOCX / CSV. Used by the applications, filings, routines and notices routers
so download plumbing is written once.

- PDF  : fpdf2 (already a dependency; mirrors the letterhead style in application_svc._build_pdf)
- DOCX : python-docx
- CSV  : stdlib csv + io.StringIO (mirrors the audit-export pattern in routers/admin_audit.py)
"""
from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field

from fastapi import HTTPException, status
from fastapi.responses import Response

ALLOWED_FORMATS = ("pdf", "docx", "csv", "xlsx")

_MEDIA_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "csv": "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

ORG_NAME = "Daffodil International University"
ORG_ADDRESS = "Daffodil Smart City, Birulia, Savar, Dhaka-1216, Bangladesh"
ORG_TAGLINE = "Generated via Anchor AI - campus.anchor-ai.bd"


@dataclass
class ExportDoc:
    """Intermediate representation rendered into PDF/DOCX/CSV."""
    title: str
    subtitle: str | None = None
    meta: list[tuple[str, str]] = field(default_factory=list)  # ordered key/value rows
    body: str | None = None                                    # free-text block
    table_headers: list[str] | None = None
    table_rows: list[list[str]] | None = None


def _pdf_safe(text: str) -> str:
    """Strip characters outside latin-1 so fpdf2 built-in fonts don't crash."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ── PDF ───────────────────────────────────────────────────────────────────────

def render_pdf(doc: ExportDoc) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(25, 25, 25)
    pdf.add_page()
    w = pdf.w - pdf.l_margin - pdf.r_margin

    # Letterhead
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(w, 10, ORG_NAME, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(w, 6, ORG_ADDRESS, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", 9)
    pdf.cell(w, 5, ORG_TAGLINE, align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(4)
    pdf.set_draw_color(11, 29, 53)
    pdf.set_line_width(0.5)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(6)

    # Title + subtitle
    pdf.set_font("Helvetica", "B", 13)
    pdf.multi_cell(w, 8, _pdf_safe(doc.title))
    if doc.subtitle:
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(w, 5, _pdf_safe(doc.subtitle))
    pdf.ln(2)

    # Meta key/value rows
    if doc.meta:
        pdf.set_font("Helvetica", "", 9)
        for k, v in doc.meta:
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(40, 6, _pdf_safe(f"{k}:"))
            pdf.set_font("Helvetica", "", 9)
            pdf.multi_cell(w - 40, 6, _pdf_safe(str(v)))
        pdf.ln(2)

    # Body
    if doc.body:
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(w, 6, _pdf_safe(doc.body))
        pdf.ln(4)

    # Table
    if doc.table_headers and doc.table_rows:
        ncol = len(doc.table_headers)
        col_w = w / ncol
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(230, 230, 230)
        for header in doc.table_headers:
            pdf.cell(col_w, 6, _pdf_safe(str(header)), border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        for row in doc.table_rows:
            for cell in row:
                pdf.cell(col_w, 6, _pdf_safe(str(cell))[:28], border=1)
            pdf.ln()

    return bytes(pdf.output())


# ── DOCX ──────────────────────────────────────────────────────────────────────

def render_docx(doc: ExportDoc) -> bytes:
    from docx import Document
    from docx.shared import Pt

    document = Document()

    document.add_heading(doc.title, level=0)
    if doc.subtitle:
        sub = document.add_paragraph()
        run = sub.add_run(doc.subtitle)
        run.italic = True
        run.font.size = Pt(10)

    for k, v in doc.meta:
        p = document.add_paragraph()
        kr = p.add_run(f"{k}: ")
        kr.bold = True
        p.add_run(str(v))

    if doc.body:
        document.add_paragraph()
        for para in doc.body.split("\n"):
            document.add_paragraph(para)

    if doc.table_headers and doc.table_rows:
        document.add_paragraph()
        table = document.add_table(rows=1, cols=len(doc.table_headers))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, header in enumerate(doc.table_headers):
            hdr[i].text = str(header)
        for row in doc.table_rows:
            cells = table.add_row().cells
            for i, cell in enumerate(row):
                cells[i].text = str(cell)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ── CSV ───────────────────────────────────────────────────────────────────────

def render_csv(doc: ExportDoc) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([doc.title])
    if doc.subtitle:
        writer.writerow([doc.subtitle])
    if doc.meta:
        writer.writerow([])
        writer.writerow(["Field", "Value"])
        for k, v in doc.meta:
            writer.writerow([k, str(v)])
    if doc.body:
        writer.writerow([])
        writer.writerow(["Body"])
        writer.writerow([doc.body])
    if doc.table_headers and doc.table_rows:
        writer.writerow([])
        writer.writerow(doc.table_headers)
        for row in doc.table_rows:
            writer.writerow([str(c) for c in row])
    return buf.getvalue().encode("utf-8-sig")  # BOM so Excel renders UTF-8 correctly


# ── XLSX ──────────────────────────────────────────────────────────────────────

def render_xlsx(doc: ExportDoc) -> bytes:
    """Render to an editable .xlsx workbook. The table block becomes a real
    header-row + data-rows sheet so it round-trips through ``parse_table``."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = "Routine"

    row = 1
    title_cell = ws.cell(row=row, column=1, value=doc.title)
    title_cell.font = Font(bold=True, size=14)
    row += 1
    if doc.subtitle:
        sub = ws.cell(row=row, column=1, value=doc.subtitle)
        sub.font = Font(italic=True, size=10)
        row += 1

    if doc.meta:
        row += 1
        for k, v in doc.meta:
            kc = ws.cell(row=row, column=1, value=str(k))
            kc.font = Font(bold=True)
            ws.cell(row=row, column=2, value=str(v))
            row += 1

    if doc.body:
        row += 1
        ws.cell(row=row, column=1, value=doc.body)
        row += 1

    if doc.table_headers and doc.table_rows:
        row += 1
        for j, header in enumerate(doc.table_headers, start=1):
            hc = ws.cell(row=row, column=j, value=str(header))
            hc.font = Font(bold=True)
        row += 1
        for data_row in doc.table_rows:
            for j, cell in enumerate(data_row, start=1):
                ws.cell(row=row, column=j, value=("" if cell is None else str(cell)))
            row += 1

        # Widen columns a little for legibility / easy editing
        for j, header in enumerate(doc.table_headers, start=1):
            width = max(12, min(40, len(str(header)) + 4))
            ws.column_dimensions[ws.cell(row=1, column=j).column_letter].width = width

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Spreadsheet parsing (re-import) ───────────────────────────────────────────

def parse_table(
    file_bytes: bytes,
    content_type: str = "",
    header_hints: tuple[str, ...] | None = None,
) -> list[dict]:
    """Parse an uploaded CSV/XLSX into a list of lowercased-header dicts.

    Standalone helper for routine re-import — independent of the timetable
    generator's own parser. Because our own exports prepend a title/meta block,
    the header row is located by ``header_hints`` (the first row containing any
    hint); without hints, the first non-empty row is the header. This lets an
    exported sheet round-trip straight back in.
    """
    ct = (content_type or "").lower()
    if "spreadsheet" in ct or "xlsx" in ct or "excel" in ct or ct.endswith(".xlsx"):
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        ws = wb.active
        raw_rows = [list(r) for r in ws.iter_rows(values_only=True)]
    else:
        text = file_bytes.decode("utf-8-sig", errors="replace")
        raw_rows = [row for row in csv.reader(io.StringIO(text))]

    # Drop fully-empty rows
    rows = [r for r in raw_rows if any(c is not None and str(c).strip() != "" for c in r)]
    if not rows:
        return []

    header_idx = 0
    if header_hints:
        hints = tuple(h.lower() for h in header_hints)
        for i, r in enumerate(rows):
            cells = [str(c).strip().lower() for c in r if c is not None]
            if any(any(h in cell for h in hints) for cell in cells):
                header_idx = i
                break

    headers = [str(h).strip().lower() if h is not None else "" for h in rows[header_idx]]
    return [
        {headers[j]: (str(v).strip() if v is not None else "") for j, v in enumerate(row) if j < len(headers)}
        for row in rows[header_idx + 1:]
    ]


# ── Dispatch ──────────────────────────────────────────────────────────────────

_RENDERERS = {"pdf": render_pdf, "docx": render_docx, "csv": render_csv, "xlsx": render_xlsx}


def _safe_filename(stem: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", stem).strip("-")
    return cleaned or "anchor-export"


def export_response(doc: ExportDoc, fmt: str, filename_stem: str) -> Response:
    """Render ``doc`` to ``fmt`` and return a download Response. 400 on bad format."""
    fmt = (fmt or "").lower()
    if fmt not in ALLOWED_FORMATS:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported format '{fmt}'. Use one of: {', '.join(ALLOWED_FORMATS)}",
        )
    content = _RENDERERS[fmt](doc)
    filename = f"{_safe_filename(filename_stem)}.{fmt}"
    return Response(
        content=content,
        media_type=_MEDIA_TYPES[fmt],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
